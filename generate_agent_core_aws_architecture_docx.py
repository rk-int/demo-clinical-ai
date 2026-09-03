import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_agent_core_aws_architecture_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Color Palette
    PRIMARY_COLOR = RGBColor(15, 23, 42)    # Slate 900 (#0f172a)
    AWS_ORANGE = RGBColor(255, 153, 0)      # AWS Orange (#ff9900)
    SECONDARY_COLOR = RGBColor(2, 132, 199) # Sky 600 (#0284c7)
    EMERALD_COLOR = RGBColor(16, 185, 129)  # Emerald 500 (#10b981)
    PURPLE_COLOR = RGBColor(124, 58, 237)   # Purple 600 (#7c3aed)
    DARK_COLOR = RGBColor(51, 65, 85)       # Slate 700 (#334155)

    def set_cell_bg(cell, hex_color):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("ENTERPRISE HEALTHCARE AI PLATFORM ON AWS")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Amazon Agent Core Architecture — Modern HIPAA-Aligned Agentic RAG Blueprint")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 1: Amazon Agent Core Migration Notice
    h1 = doc.add_heading("1. Strategic Update: Transition to Amazon Agent Core", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph(
        "With legacy/classic agent orchestration frameworks sunset, this architecture fully integrates "
        "Amazon Agent Core—the next-generation AWS agentic platform engine. Amazon Agent Core provides "
        "native multi-agent collaboration, built-in agentic memory management, secure tool calling, and automated "
        "clinical evaluation harnesses required for healthcare production deployments."
    )
    p1.runs[0].font.size = Pt(10.5)

    # Section 2: Comprehensive 10-Domain AWS Tech Stack Mapping Table with Amazon Agent Core
    h2 = doc.add_heading("2. Comprehensive Feature & Amazon Agent Core Tech Stack Table", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    t_map = doc.add_table(rows=1, cols=4)
    t_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_map.rows[0].cells
    m_hdr[0].text = "Architecture Domain"
    m_hdr[1].text = "Current Local Prototype"
    m_hdr[2].text = "Target AWS Tech Stack (Amazon Agent Core Architecture)"
    m_hdr[3].text = "Architectural Purpose & Technical Role"

    for cell in m_hdr:
        set_cell_bg(cell, "0F172A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    master_mapping = [
        # Domain 1: Edge & Ingress
        (
            "1. User Channels, Secure Edge & Gateway",
            "React 18 SPA, Vite dev server (`npm run dev`), simulated TLS badge",
            "AWS Amplify, Amazon CloudFront, Amazon Route 53, AWS WAF, AWS Shield Advanced, Amazon API Gateway",
            "Global low-latency edge delivery via CloudFront CDN, Route 53 DNS failover, AWS WAF/Shield L7 DDoS protection, API Gateway zero-trust mTLS ingress with rate limiting."
        ),
        # Domain 2: Voice & Multimodal AI
        (
            "2. Clinical Voice & Ambient AI Services",
            "Text-only prompt input in `KnowledgeQAView.tsx`",
            "Amazon Transcribe Medical, Amazon Lex, Amazon Textract, Comprehend Medical, Amazon Polly, Amazon Translate",
            "Speech-to-text for clinical voice dictation (Transcribe Medical), conversational voice bots (Lex), medical OCR chart extraction (Textract), and multi-lingual translation."
        ),
        # Domain 3: Identity, Access & SSO
        (
            "3. Identity, Access & SSO",
            "Simulated JWT state (`auth.types.ts`) with clinician/admin roles",
            "Amazon Cognito, AWS IAM / IAM Identity Center, AWS Managed SSO",
            "User Pools for clinician login, SAML 2.0 / OIDC federation for hospital EMR SSO (Epic/Cerner IDP), custom JWT role claims (`isAdmin`), and fine-grained IAM policy scoping."
        ),
        # Domain 4: Agentic Orchestration (AMAZON AGENT CORE)
        (
            "4. Agentic Orchestration & Agent Core",
            "State machine state (`VerticalPatientSearchFlowCanvas.tsx`, `agentEngine.ts`)",
            "Amazon Agent Core (Multi-Agent Collaboration Engine), Amazon Bedrock, AWS Lambda, AWS ECS Fargate",
            "Amazon Agent Core orchestrates multi-agent delegation across Orchestrator, Clinical, Workflow, and Compliance agents. Lambda executes serverless tool actions (FHIR queries, order drafting)."
        ),
        # Domain 5: Context & Memory (AMAZON AGENT CORE MEMORY)
        (
            "5. Context, Session & Agent Core Memory",
            "React in-memory component state",
            "Amazon Agent Core Memory, Amazon ElastiCache for Redis, Amazon DynamoDB",
            "Amazon Agent Core Memory maintains long-term clinical conversation context, patient session state, and multi-agent context fusion. ElastiCache Redis handles LLM semantic caching."
        ),
        # Domain 6: Foundation Models & Guardrails
        (
            "6. Foundation Model Fabric & Guardrails",
            "Simulated LLM response generator in `agentEngine.ts`",
            "Amazon Bedrock (Claude 3.5 Sonnet), SageMaker Endpoints, Amazon Bedrock Guardrails",
            "Bedrock managed access to Anthropic Claude 3.5 Sonnet. SageMaker for fine-tuned clinical models. Bedrock Guardrails for DLP PHI redaction, prompt injection defense, & zero-hallucination checks."
        ),
        # Domain 7: Advanced RAG & Knowledge Graph
        (
            "7. Advanced RAG & Knowledge Graph",
            "Static guideline arrays in `KnowledgeQAView.tsx`",
            "Amazon Bedrock Knowledge Bases, Amazon OpenSearch Service, Amazon Neptune (Health KG)",
            "Bedrock Knowledge Bases parses clinical PDFs. OpenSearch Service handles hybrid vector + BM25 search. Amazon Neptune manages Health Knowledge Graph (Care Pathways & Entity Linking)."
        ),
        # Domain 8: Data Platform & EHR Integration
        (
            "8. Governed Data Platform & Integration",
            "Synthetic FHIR R4 JSON schemas (`syntheticFhirData.ts`)",
            "AWS HealthLake, Amazon S3 Lakehouse, AWS Lake Formation, Amazon Aurora (RDS), Amazon Redshift, AWS AppFlow, Amazon MSK, Amazon MQ",
            "AWS HealthLake manages HIPAA FHIR R4 store. S3 Lakehouse + Glue + Athena for analytics. Lake Formation for FGAC governance. AppFlow/MSK/MQ for HL7v2, DICOM, and streaming EHR events."
        ),
        # Domain 9: Observability & Agent Core Evaluation
        (
            "9. Observability & Agent Core Evaluation",
            "Console logs & step state counters (`currentExecutionStep`)",
            "Amazon Agent Core Evaluation Harness, Amazon CloudWatch, AWS X-Ray, AWS CodePipeline",
            "Amazon Agent Core Evaluation Harness measures agent task accuracy, RAG faithfulness, recall, & zero-hallucination metrics. X-Ray provides end-to-end multi-agent latency tracing."
        ),
        # Domain 10: Security Overlay & Multi-Region Resilience
        (
            "10. Security Overlay & Multi-Region Resilience",
            "Local environment variables and in-memory execution",
            "AWS KMS, AWS CloudHSM, AWS Secrets Manager, Amazon Macie, Amazon GuardDuty, AWS Security Hub, AWS Config, AWS CloudTrail, AWS Global Accelerator, Aurora Global Database, AWS Backup",
            "KMS envelope encryption, CloudHSM FIPS 140-2 L3 key vault, Macie automated PHI discovery in S3, GuardDuty threat detection, CloudTrail immutable audit logs, Global Accelerator + Aurora Global DB for multi-region failover."
        )
    ]

    for domain, local_impl, aws_svc, role_desc in master_mapping:
        row_cells = t_map.add_row().cells
        row_cells[0].text = domain
        row_cells[1].text = local_impl
        row_cells[2].text = aws_svc
        row_cells[3].text = role_desc
        set_cell_bg(row_cells[0], "F8FAFC")
        set_cell_bg(row_cells[2], "E0F2FE") # Sky Blue highlight for Amazon Agent Core & AWS Services

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Detailed Breakdown of Amazon Agent Core Features
    h3 = doc.add_heading("3. Key Architectural Roles of Amazon Agent Core", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    agent_core_highlights = [
        ("Amazon Agent Core Collaboration Engine", "Native multi-agent orchestration replacing legacy agent code. Dynamically routes clinical queries from Orchestrator Agent to Specialist Knowledge, Patient EHR, and Workflow agents."),
        ("Amazon Agent Core Memory", "Built-in long-term & short-term agent memory management. Preserves clinical conversation context, patient timeline state, and multi-modal payload synthesis across multi-turn sessions."),
        ("Amazon Agent Core Tool Gateway", "Secure, zero-trust execution sandbox for agent tool calls (e.g. querying AWS HealthLake FHIR resources, scheduling appointments, or drafting specialist referrals)."),
        ("Amazon Agent Core Evaluation Harness", "Automated LLMOps testing suite measuring agent task completion accuracy, RAG groundedness, hallucination rates, and clinical safety compliance before production release.")
    ]

    for title, desc in agent_core_highlights:
        p_ac = doc.add_paragraph()
        p_ac.paragraph_format.left_indent = Inches(0.2)
        p_ac.paragraph_format.space_after = Pt(4)

        r_title = p_ac.add_run(f"⚡ {title}: ")
        r_title.font.bold = True
        r_title.font.color.rgb = PURPLE_COLOR

        r_desc = p_ac.add_run(desc)
        r_desc.font.size = Pt(10)
        r_desc.font.color.rgb = DARK_COLOR

    # Save Document
    doc_path = "/Users/rk/Antigravity/demo/demo-clinical-ai/Clinical_AI_Portal_AWS_Architecture_Blueprint.docx"
    doc.save(doc_path)
    print(f"Successfully generated Amazon Agent Core AWS Architecture Word document at: {doc_path}")

if __name__ == "__main__":
    create_agent_core_aws_architecture_docx()
