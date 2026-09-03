import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_aws_architecture_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Tokens
    PRIMARY_COLOR = RGBColor(15, 23, 42)    # Slate 900 (#0f172a)
    AWS_ORANGE = RGBColor(255, 153, 0)      # AWS Orange (#ff9900)
    SECONDARY_COLOR = RGBColor(2, 132, 199) # Sky 600 (#0284c7)
    EMERALD_COLOR = RGBColor(16, 185, 129)  # Emerald 500 (#10b981)
    DARK_COLOR = RGBColor(51, 65, 85)       # Slate 700 (#334155)

    def set_cell_bg(cell, hex_color):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("HealthNet Clinical AI Portal — End-to-End AWS Production Architecture Blueprint")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Feature-by-Feature Mapping from Local Prototype to Production AWS Services")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Section 1: Executive Overview
    h1 = doc.add_heading("1. Executive Summary & Production Blueprint Goal", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph(
        "This document provides a comprehensive, component-by-component architectural blueprint for migrating "
        "the HealthNet Clinical AI Portal from its current local React/Express prototype into a 100% cloud-native, "
        "HIPAA-eligible, multi-tenant enterprise architecture deployed entirely on Amazon Web Services (AWS)."
    )
    p1.runs[0].font.size = Pt(10.5)

    # Section 2: Complete Feature-to-AWS Service Mapping Table
    h2 = doc.add_heading("2. Complete Feature & AWS Tech Stack Mapping Table", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    t_map = doc.add_table(rows=1, cols=4)
    t_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_map.rows[0].cells
    m_hdr[0].text = "Portal Module / Feature"
    m_hdr[1].text = "Current Local Prototype Implementation"
    m_hdr[2].text = "Target AWS Tech Stack Service"
    m_hdr[3].text = "Production Architectural Role & AWS Capabilities"

    for cell in m_hdr:
        set_cell_bg(cell, "0F172A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    mapping_data = [
        (
            "Frontend Hosting & Edge CDN",
            "React 18, TypeScript, Vite local dev server (`npm run dev`)",
            "AWS Amplify / Amazon CloudFront + Amazon S3",
            "S3 bucket hosting static bundle assets, CloudFront CDN for global low-latency SSL/TLS distribution, Amplify for automated Git CI/CD deployments."
        ),
        (
            "Zero-Trust API Gateway & Ingress",
            "Express / Node.js local proxy (`server.ts`) with TLS 1.3 badge",
            "Amazon API Gateway + AWS WAF",
            "Manages mTLS endpoint security, CORS headers, rate limiting, request validation, and routes API requests safely to backend Lambda/ECS microservices with WAF protection."
        ),
        (
            "Authentication & RBAC Security",
            "Simulated OAuth2/JWT auth token state with `PORTAL_ADMIN` & `CLINICIAN` roles",
            "Amazon Cognito + AWS IAM",
            "User Pools for patient & clinician identity, OAuth2 / OIDC SAML federation for hospital SSO (Epic/Cerner IDP), custom JWT role claims (`isAdmin`), and IAM fine-grained session tokens."
        ),
        (
            "Multi-Tenant Hospital Federation",
            "Global `HospitalNetworkContext` filtering 3 hospital facilities",
            "Amazon DynamoDB / Amazon Aurora PostgreSQL (RLS)",
            "Multi-tenant hospital configuration tables with Row-Level Security (RLS) policies in Aurora PostgreSQL ensuring strict data isolation per hospital facility tenant."
        ),
        (
            "Patient 360 & Healthcare Data Store",
            "Synthetic FHIR R4 JSON data (`syntheticFhirData.ts`, `patient.types.ts`)",
            "AWS HealthLake + Amazon S3 Data Lake",
            "HIPAA-eligible managed FHIR R4 store with integrated medical NLP. S3 Data Lake (Parquet/Iceberg) + Amazon Athena for ad-hoc clinical data analytics."
        ),
        (
            "Multi-Agent AI Orchestration Engine",
            "Simulated state machine (`VerticalPatientSearchFlowCanvas.tsx`, `agentEngine.ts`)",
            "Amazon Bedrock Agents + AWS Lambda / AWS ECS Fargate",
            "Bedrock Agents orchestrate task decomposition, routing requests across specialized agent functions. Lambda executes serverless tool actions (FHIR queries, appointment bookings)."
        ),
        (
            "Foundation LLM Models & Reasoning",
            "Simulated Gemini 1.5 Pro / Claude response generation",
            "Amazon Bedrock (Claude 3.5 Sonnet / Claude 3 Opus)",
            "Managed serverless access to state-of-the-art foundation models (Anthropic Claude 3.5 Sonnet) providing high-precision clinical reasoning without managing infrastructure."
        ),
        (
            "Clinical RAG Knowledge Base",
            "Static clinical guideline arrays (KDIGO, GOLD, AHA) in `KnowledgeQAView.tsx`",
            "Amazon Bedrock Knowledge Bases + Amazon OpenSearch Serverless",
            "Automated RAG pipeline parsing medical PDFs/guidelines, generating vector embeddings via Amazon Titan Multimodal Embeddings, stored in OpenSearch Serverless for hybrid vector search."
        ),
        (
            "Clinical Guardrails, DLP & PHI Redaction",
            "Regex DLP PHI redaction, zero-hallucination token check, citation validator",
            "Amazon Bedrock Guardrails + Amazon Comprehend Medical",
            "Bedrock Guardrails enforces real-time PII/PHI redaction, prompt injection protection, hallucination evaluation, and deny topics. Comprehend Medical extracts ICD-10/RxNorm codes."
        ),
        (
            "Appointments & Encounter Gateway",
            "Lightweight booking modal, simulated Email & SMS receipts, close encounter modal",
            "Amazon SNS + Amazon SES + AWS Lambda + Amazon Aurora",
            "Amazon SNS dispatches transactional SMS alerts to patient phones. Amazon SES sends HTML email confirmations. Lambda handles booking logic and updates Aurora PostgreSQL state."
        ),
        (
            "Executive Operations & ROI Dashboard",
            "React SVG charts for Value Realized, AI Precision line graph, cost avoidance cards",
            "Amazon QuickSight + Amazon Redshift / Aurora Serverless v2",
            "Serverless BI dashboards with embedded QuickSight analytics, automated ML forecasting for patient throughput, and Redshift serverless data warehouse for executive analytics."
        ),
        (
            "HIPAA Audit & Compliance Center",
            "React stream component rendering security logs, PHI redactions, security score",
            "AWS CloudTrail + Amazon CloudWatch Logs + AWS Security Hub",
            "CloudTrail maintains immutable, tamper-proof governance audit logs. CloudWatch Logs captures encrypted agent logs with metric alarms. Security Hub monitors HIPAA compliance."
        ),
        (
            "Data Encryption & Key Management",
            "Local environment variables and in-memory encryption",
            "AWS KMS + AWS Secrets Manager",
            "AWS KMS manages Customer Managed Keys (CMK) for Envelope Encryption (AES-256) at rest across S3, Aurora, HealthLake, and DynamoDB. Secrets Manager handles API key rotation."
        ),
        (
            "Network Privacy & Virtual Isolation",
            "Localhost network port binding (`localhost:3000`)",
            "Amazon VPC + AWS PrivateLink + VPC Endpoints",
            "Isolated Virtual Private Cloud (VPC) with private subnets for microservices and database. AWS PrivateLink connects securely to Bedrock and HealthLake without traversing public internet."
        ),
        (
            "Distributed Agent Tracing & Monitoring",
            "Console logging and step state counters (`currentExecutionStep`)",
            "AWS X-Ray + Amazon CloudWatch ServiceLens",
            "Provides end-to-end latency tracing across multi-agent microservice calls, mapping exact execution timing from Gateway -> Orchestrator -> Bedrock LLM -> HealthLake."
        ),
        (
            "Infrastructure as Code (IaC) & Deployment",
            "Manual file editing and Vite build scripts",
            "AWS CDK (TypeScript) / AWS CloudFormation",
            "Declarative TypeScript CDK scripts defining the complete infrastructure topology as code, enabling repeatable environment provisioning (Dev, Staging, Production)."
        )
    ]

    for module, local_impl, aws_svc, role_desc in mapping_data:
        row_cells = t_map.add_row().cells
        row_cells[0].text = module
        row_cells[1].text = local_impl
        row_cells[2].text = aws_svc
        row_cells[3].text = role_desc
        set_cell_bg(row_cells[0], "F8FAFC")
        set_cell_bg(row_cells[2], "FEF3C7") # Light Amber for AWS Services

    # Save Document
    doc_path = "/Users/rk/Antigravity/demo/demo-clinical-ai/Clinical_AI_Portal_AWS_Architecture_Blueprint.docx"
    doc.save(doc_path)
    print(f"Successfully generated AWS Architecture Blueprint Word document at: {doc_path}")

if __name__ == "__main__":
    create_aws_architecture_docx()
