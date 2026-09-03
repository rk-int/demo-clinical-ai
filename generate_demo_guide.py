import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_demo_guide_docx():
    doc = docx.Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style Helpers
    PRIMARY_COLOR = RGBColor(12, 74, 110)    # Deep Cyan / Teal (#0c4a6e)
    SECONDARY_COLOR = RGBColor(16, 185, 129) # Emerald Green (#10b981)
    DARK_COLOR = RGBColor(30, 41, 59)        # Slate 800 (#1e293b)
    LIGHT_BG = "F0FDF4"                       # Light Emerald Tint
    BOX_BG = "F8FAFC"                         # Slate 50

    # Helper function for cell background color
    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("HealthNet Clinical AI Assistant & Enterprise Portal")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Step-by-Step Demo Script, Technical Architecture & AWS Migration Plan")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Important Prototype Disclaimer Box
    table_disc = doc.add_table(rows=1, cols=1)
    table_disc.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_disc = table_disc.cell(0, 0)
    set_cell_background(cell_disc, "FEF2F2") # Light Red Tint
    
    p_disc = cell_disc.paragraphs[0]
    p_disc.paragraph_format.space_before = Pt(8)
    p_disc.paragraph_format.space_after = Pt(8)
    p_disc.paragraph_format.left_indent = Inches(0.2)
    p_disc.paragraph_format.right_indent = Inches(0.2)
    
    r_disc_title = p_disc.add_run("⚠️ MANDATORY PRESENTER OPENING DISCLAIMER\n")
    r_disc_title.font.bold = True
    r_disc_title.font.size = Pt(11)
    r_disc_title.font.color.rgb = RGBColor(185, 28, 28)
    
    r_disc_text = p_disc.add_run(
        "\"Good day everyone. Before we begin, I would like to state clearly that this application is an interactive PROTOTYPE and DEMO CONCEPT. "
        "It is designed to demonstrate end-to-end clinical AI workflows, multi-agent orchestration, and clinician user experience. "
        "Currently, all AI models, agent workflows, synthetic FHIR data stores, and guardrail validations are running on a LOCAL technology stack. "
        "This is not yet a production deployment. We have a detailed AWS Cloud migration architecture designed to transition this prototype into a production-grade AWS environment.\""
    )
    r_disc_text.font.size = Pt(10.5)
    r_disc_text.font.italic = True
    r_disc_text.font.color.rgb = DARK_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: Executive Summary & Overview
    h1 = doc.add_heading("1. Executive Summary & Application Overview", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph(
        "The HealthNet Clinical AI Portal is a next-generation federated clinical decision support platform. "
        "It enables healthcare organizations across multi-hospital networks to leverage artificial intelligence for patient chart synthesis, "
        "clinical guideline lookup, automated appointment management, executive analytics, and zero-trust HIPAA-compliant auditability."
    )
    p1.runs[0].font.size = Pt(11)

    # Section 2: Technical Stack (Local Prototype)
    h2 = doc.add_heading("2. Local Technology Stack (Prototype Architecture)", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    tech_data = [
        ("Frontend UI Framework", "React 18 with TypeScript & Vite (Ultra-fast HMR)"),
        ("Styling & Graphics", "Vanilla CSS Design Tokens, Tailwind CSS, Custom SVG Motion Graphics"),
        ("Icons & UI Components", "Lucide React Icons, Glassmorphism Cards, Dynamic Theme Engine"),
        ("Local Express Proxy Server", "Node.js / Express proxy (`server.ts`) for endpoint handling"),
        ("Healthcare Data Standard", "FHIR R4 (Fast Healthcare Interoperability Resources) Synthetic Data"),
        ("AI Orchestration Architecture", "Simulated Multi-Agent System (Knowledge Agent, Patient EHR Agent, Workflow Agent)"),
        ("Guardrails & Safety Engine", "6-Stage Context Fusion & LLM Guardrails (DLP PHI Redaction, Zero-Hallucination Checks)")
    ]

    t_tech = doc.add_table(rows=1, cols=2)
    t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t_tech.rows[0].cells
    hdr_cells[0].text = "Component Layer"
    hdr_cells[1].text = "Technology / Implementation Details"
    for cell in hdr_cells:
        set_cell_background(cell, "0F172A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    for row_item in tech_data:
        row_cells = t_tech.add_row().cells
        row_cells[0].text = row_item[0]
        row_cells[1].text = row_item[1]
        set_cell_background(row_cells[0], "F1F5F9")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Step-by-Step, Screen-by-Screen Demo Script
    h3 = doc.add_heading("3. Step-by-Step Screen Demo Script (10-Minute Presentation)", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    steps = [
        {
            "num": "Screen 1",
            "name": "Landing Page & Role Selection Modal",
            "action": "Show landing page. Click 'Sign In / Select Portal Role'. Select 'Attending Clinician (Dr. Sarah Chen)' or 'Portal Admin'.",
            "say": "\"Welcome to the HealthNet Portal. Our login system features Role-Based Access Control (RBAC). Clinicians get clinical AI decision support, while Portal Admins get security audit tools.\"",
            "tech": "Simulated JWT/OAuth authentication token state, setting role permissions (`PORTAL_ADMIN` vs `CLINICIAN`)."
        },
        {
            "num": "Screen 2",
            "name": "Multi-Hospital Workspace & Network Selector",
            "action": "Point to top navigation bar. Click the Hospital Selector dropdown and switch between 'HealthNet Central Network', 'St. Jude Medical Center', and 'Metro General Hospital'.",
            "say": "\"Notice the top bar. HealthNet supports a Multi-Hospital Federated Network. When I switch from St. Jude to Metro General, all patient records, metrics, and agent search scopes dynamically update instantly without page reloads.\"",
            "tech": "Global React Context (`HospitalNetworkContext`) dynamically filters FHIR dataset queries and workspace statistics."
        },
        {
            "num": "Screen 3",
            "name": "Patient 360 View & Search",
            "action": "Click 'Patients' in sidebar. Select patient 'Eleanor Vance'. Review Vitals, Timeline, and Attending Doctor thumbnail.",
            "say": "\"Here is the Patient 360 view. It aggregates synthetic FHIR R4 records including vital signs, active diagnoses, lab observations, and attending clinician photo lookup (e.g. Dr. Emily Vance, MD).\"",
            "tech": "Parses FHIR Patient, Observation, and Encounter resources into unified 360-degree timeline cards."
        },
        {
            "num": "Screen 4",
            "name": "Clinical AI Assistant & Integrated Live Agent Flow (⭐ Highlight)",
            "action": "Click 'AI Assistant' tab. Select sample query ('Evaluate Eleanor Vance for CKD Stage 3'). Click 'Submit'. Click 'Expand Live Agent Flow'.",
            "say": "\"This is our core innovation: The Clinical AI Assistant with Live Agent Execution Flow. Notice how the Orchestrator dispatches the request, and a moving blue dashed arrow points ONLY to the triggered Knowledge Agent, while standby agents stay clean. Watch the 10-stage sequential pipeline flow down into Context Fusion, Prompt Assembly, Pre-Guardrail DLP, Gemini LLM reasoning, Post-Guardrail zero-hallucination check, and terminates cleanly at Response Validation!\"",
            "tech": "10-Stage sequential execution state machine with animated SVG paths (`animate-moving-arrow-dash`), DLP/PHI redaction regex, and zero-hallucination token validation."
        },
        {
            "num": "Screen 5",
            "name": "Simplified Appointments Center",
            "action": "Click 'Appointments' in sidebar. Click 'Book New Encounter'. Fill basic contact info (Name, Age, Phone, Email) and submit.",
            "say": "\"In the Appointments tab, clinicians or staff can quickly schedule patient encounters without creating bloated dummy records. Automated SMS and email confirmation receipts are simulated upon booking.\"",
            "tech": "Creates lightweight appointment object and simulates asynchronous twilio/sendgrid SMS and Email dispatch."
        },
        {
            "num": "Screen 6",
            "name": "Executive Dashboard (ROI & AI Performance)",
            "action": "Click 'Executive Dashboard' in sidebar.",
            "say": "\"The Executive Dashboard provides system-wide visibility for all stakeholders. Key metrics show 8,542 patients impacted, 1,248 documentation hours saved, 62% automation rate, 0 safety incidents, and $342,800 in estimated cost avoidance. Interactive SVG charts show Value Realized and AI Model Precision trends.\"",
            "tech": "Aggregates operational KPIs, responsive SVG bar charts and multi-series line graphs filtered by hospital scope."
        },
        {
            "num": "Screen 7",
            "name": "Audit & Compliance Center (Admin Only)",
            "action": "Switch user role to 'Portal Admin'. Click 'Audit Center' in sidebar.",
            "say": "\"For strict security compliance, the Audit Center is restricted to Portal Admins. It displays real-time HIPAA access logs, PHI redaction events, and security readiness scores filtered dynamically by hospital.\"",
            "tech": "RBAC route guard (`isAdmin` restriction) rendering real-time security log streams and compliance badges."
        }
    ]

    for step in steps:
        p_step_title = doc.add_paragraph()
        r_num = p_step_title.add_run(f"📌 {step['num']}: {step['name']}\n")
        r_num.font.bold = True
        r_num.font.size = Pt(12)
        r_num.font.color.rgb = PRIMARY_COLOR

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.2)
        
        r_act = p_desc.add_run("👉 What to Click/Show: ")
        r_act.font.bold = True
        p_desc.add_run(f"{step['action']}\n\n")

        r_say = p_desc.add_run("💬 What to Say (Script): ")
        r_say.font.bold = True
        r_say_text = p_desc.add_run(f"{step['say']}\n\n")
        r_say_text.font.italic = True
        r_say_text.font.color.rgb = RGBColor(30, 58, 138)

        r_tch = p_desc.add_run("⚙️ Under the Hood (Technical): ")
        r_tch.font.bold = True
        p_desc.add_run(f"{step['tech']}")

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Section 4: In-Scope vs Out-of-Scope
    h4 = doc.add_heading("4. Scope Boundary (In-Scope vs. Out-of-Scope)", level=1)
    h4.runs[0].font.color.rgb = PRIMARY_COLOR

    t_scope = doc.add_table(rows=1, cols=2)
    t_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_hdr = t_scope.rows[0].cells
    s_hdr[0].text = "In-Scope (Current Prototype)"
    s_hdr[1].text = "Out-of-Scope (Future AWS Production)"
    for cell in s_hdr:
        set_cell_background(cell, "0C4A6E")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    scope_matrix = [
        ("Full Interactive Web UI/UX with rich aesthetics", "Live EMR integration with Epic / Cerner webhooks"),
        ("Federated Multi-Hospital switching context", "Real patient PHI handling (only synthetic FHIR used)"),
        ("Integrated 10-step live agent execution flow", "Multi-region AWS cloud deployment"),
        ("Role-Based Access Control (Admin vs Clinician)", "Live production LLM API billing keys"),
        ("Executive KPI Dashboard & ROI Analytics", "Physical SMS/Email cellular gateway hardware"),
        ("Admin HIPAA Audit Trail & Security View", "HL7 v2 legacy engine integration")
    ]

    for in_s, out_s in scope_matrix:
        row_cells = t_scope.add_row().cells
        row_cells[0].text = f"✅ {in_s}"
        row_cells[1].text = f"⏳ {out_s}"

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 5: Future Implementation & AWS Cloud Migration Architecture
    h5 = doc.add_heading("5. Production AWS Cloud Migration Plan", level=1)
    h5.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph(
        "To transition this local prototype into a high-availability, HIPAA-compliant enterprise platform, "
        "we have designed a comprehensive AWS Cloud Migration Architecture mapping each prototype component to native AWS services:"
    )

    aws_mapping = [
        ("Frontend Hosting & CDN", "AWS Amplify / Amazon CloudFront + S3", "Global low-latency HTTPS delivery, automated CI/CD deployment, edge caching."),
        ("API Gateway & Zero-Trust", "Amazon API Gateway", "Enforces mTLS, rate limiting, CORS, and routes traffic securely to backend services."),
        ("User Authentication & RBAC", "Amazon Cognito", "HIPAA-compliant Identity Provider supporting SAML/OAuth2, MFA, and role claims."),
        ("Agent Microservices Runtime", "AWS ECS (Fargate) / EKS", "Serverless containerized execution for Orchestrator and Specialist Agents."),
        ("LLM & Generative AI Engine", "Amazon Bedrock (Claude 3.5 / Titan)", "Managed access to top clinical LLMs with built-in Bedrock Guardrails for PHI redaction."),
        ("Healthcare Data Store", "AWS HealthLake", "HIPAA-eligible FHIR R4 datastore with automated clinical entity extraction."),
        ("Relational Database", "Amazon Aurora PostgreSQL Serverless v2", "Encrypted multi-AZ relational database for user preferences, audit trails, and schedules."),
        ("Audit & Compliance Logging", "AWS CloudTrail & Amazon CloudWatch", "Immutable security logging, real-time threat detection, and HIPAA compliance reporting."),
        ("Distributed Agent Tracing", "AWS X-Ray", "Provides end-to-end latency tracing across multi-agent microservice calls."),
        ("Data Security & Encryption", "AWS KMS & AWS WAF", "Customer Managed Keys for Envelope Encryption at rest and Web Application Firewall protection.")
    ]

    t_aws = doc.add_table(rows=1, cols=3)
    t_aws.alignment = WD_TABLE_ALIGNMENT.CENTER
    a_hdr = t_aws.rows[0].cells
    a_hdr[0].text = "Prototype Layer"
    a_hdr[1].text = "Target AWS Service"
    a_hdr[2].text = "Production Architectural Role"
    for cell in a_hdr:
        set_cell_background(cell, "065F46") # Emerald 800
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    for p_layer, aws_svc, role_desc in aws_mapping:
        row_cells = t_aws.add_row().cells
        row_cells[0].text = p_layer
        row_cells[1].text = aws_svc
        row_cells[2].text = role_desc
        set_cell_background(row_cells[0], "F0FDF4")

    # Save document
    doc_path = "/Users/rk/Antigravity/demo/demo-clinical-ai/Clinical_AI_Portal_Demo_Guide.docx"
    doc.save(doc_path)
    print(f"Successfully generated editable Word document at: {doc_path}")

if __name__ == "__main__":
    create_demo_guide_docx()
