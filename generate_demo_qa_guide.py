import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_demo_qa_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Tokens
    PRIMARY_COLOR = RGBColor(15, 23, 42)    # Slate 900 (#0f172a)
    SECONDARY_COLOR = RGBColor(2, 132, 199) # Sky 600 (#0284c7)
    EMERALD_COLOR = RGBColor(16, 185, 129)  # Emerald 500 (#10b981)
    DARK_COLOR = RGBColor(51, 65, 85)       # Slate 700 (#334155)

    def set_cell_bg(cell, hex_color):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("HealthNet Clinical AI Portal — Stakeholder Q&A & Feature Capability Guide")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Anticipated Demo Questions, Supported vs. Future Capabilities & AWS Cloud Roadmap")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Section 1: Executive Overview
    h1 = doc.add_heading("1. Executive Overview & Demo Briefing Context", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph(
        "This document equips the presenter with direct, authoritative answers to high-frequency technical, clinical, "
        "and architectural questions anticipated during today's live demonstration of the HealthNet Clinical AI Portal. "
        "It details current functional capabilities, boundaries of the local prototype, and the enterprise AWS production roadmap."
    )
    p1.runs[0].font.size = Pt(10.5)

    # Section 2: Anticipated Demo Questions & Authoritative Answers
    h2 = doc.add_heading("2. Anticipated Demo Q&A (Questions & Technical Answers)", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    qa_list = [
        {
            "category": "Architecture & EHR Integration",
            "q": "Q1: Is this portal connected live to our production EHR systems (e.g. Epic, Cerner, Allscripts)?",
            "a": "Currently, NO. This application is an interactive local prototype powered by synthetic FHIR R4 JSON schemas (Patients, Encounters, Observations, Diagnostic Reports). "
                 "In production, the platform will connect via AWS HealthLake and HL7 FHIR R4 APIs using secure SMART-on-FHIR OAuth2 adapters."
        },
        {
            "category": "AI Safety & Zero-Hallucination",
            "q": "Q2: How does the AI Assistant ensure clinical accuracy and prevent hallucinations?",
            "a": "The assistant executes a 10-stage integrated execution pipeline. Before generating output, Stage 7 (Pre-Guardrail) enforces DLP/PHI redaction. "
                 "Following LLM inference, Stage 9 (Post-Guardrail) performs zero-hallucination token validation against ground-truth medical documents, and Stage 10 (Response Validation) validates LOINC/ICD-10 code mappings and medical citations before rendering."
        },
        {
            "category": "Multi-Hospital Scalability",
            "q": "Q3: How does the multi-hospital network selection work under the hood?",
            "a": "The portal uses a global React Context (`HospitalNetworkContext`) managing multi-tenant hospital scopes (e.g., St. Jude, Metro General, HealthNet Central). "
                 "Selecting a hospital dynamically filters synthetic patient cohorts, clinical metrics, audit trail logs, and agent execution scopes in real time without refreshing the web page."
        },
        {
            "category": "Security & HIPAA Compliance",
            "q": "Q4: How is patient PHI protected and who can see security audit logs?",
            "a": "Patient PHI is protected via Zero-Trust TLS 1.3 ingress at the AI Gateway and regex-based DLP redaction. "
                 "Access Control uses Role-Based Access Control (RBAC). Clinical tools are available to Clinicians, whereas the Audit & Compliance Center is strictly restricted to Portal Admins via RBAC route guards (`isAdmin`)."
        },
        {
            "category": "Appointments & Encounter Management",
            "q": "Q5: How does the Appointment Booking and Closure workflow function?",
            "a": "Clinicians or staff click 'Book Appointment' to launch a lightweight modal requiring only basic patient details (Name, Age, Phone, Email, Facility, Department, Date, Time, Reason). "
                 "Upon booking, it dispatches simulated automated Email & SMS alerts and appends the record to the Scheduled Appointments Register. Clinicians can also click 'Close Appointment' to record encounter completion, capture closure reasons, and send completion receipts."
        },
        {
            "category": "Executive ROI & Performance",
            "q": "Q6: What operational metrics and ROI figures does the Executive Dashboard track?",
            "a": "The Executive Dashboard tracks 5 core KPIs: 8,542 Patients Impacted, 1,248 Documentation Hours Saved (▲ 18.9%), -40% Chart Review Time, 62% Automation Rate, and 0 Safety Incidents. "
                 "It also displays estimated cost avoidance ($342,800 saved) and interactive SVG charts for Value Realized and AI Model Precision trends."
        },
        {
            "category": "AWS Production Deployment",
            "q": "Q7: How will this local prototype be migrated to AWS Cloud for production deployment?",
            "a": "We have a full AWS deployment architecture: Frontend hosted on AWS Amplify/CloudFront, API Gateway for mTLS zero-trust routing, Amazon Cognito for SAML/OAuth2 auth, AWS ECS (Fargate) for multi-agent container microservices, Amazon Bedrock for managed clinical LLMs with Bedrock Guardrails, AWS HealthLake for FHIR datastores, Amazon Aurora Serverless v2 for database, and AWS CloudTrail/X-Ray for HIPAA audit trails."
        }
    ]

    for item in qa_list:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)

        r_cat = p_q.add_run(f"[{item['category']}]\n")
        r_cat.font.bold = True
        r_cat.font.size = Pt(9.5)
        r_cat.font.color.rgb = SECONDARY_COLOR

        r_q = p_q.add_run(item['q'])
        r_q.font.bold = True
        r_q.font.size = Pt(11)
        r_q.font.color.rgb = PRIMARY_COLOR

        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.paragraph_format.space_after = Pt(8)

        r_lbl = p_a.add_run("Answer: ")
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = EMERALD_COLOR

        r_txt = p_a.add_run(item['a'])
        r_txt.font.size = Pt(10)
        r_txt.font.color.rgb = DARK_COLOR

    # Section 3: Feature Matrix (Supported vs. Yet to be Supported)
    h3 = doc.add_heading("3. Feature Capability Matrix (Supported vs. Future Roadmap)", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    t_feat = doc.add_table(rows=1, cols=3)
    t_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_hdr = t_feat.rows[0].cells
    f_hdr[0].text = "Portal Module / Area"
    f_hdr[1].text = "Currently Supported (Demo Prototype Scope)"
    f_hdr[2].text = "Yet to be Supported (AWS Production Roadmap)"
    for cell in f_hdr:
        set_cell_bg(cell, "0F172A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    matrix = [
        (
            "Multi-Hospital Federation",
            "Dynamic hospital context switching across 3 facility networks with instant UI/dataset filtering.",
            "Multi-region data residency controls & cross-border compliance routing."
        ),
        (
            "Patient 360 View",
            "Timeline visualization, vitals, active diagnoses, lab trends, and clinician photo lookup.",
            "Real-time streaming telemetry from ICU monitors & wearable device integrations."
        ),
        (
            "Clinical AI Assistant",
            "10-stage execution pipeline, single-agent SVG arrow routing, zero-shot prompt assembly, DLP PHI redaction.",
            "Live multi-modal DICOM medical image analysis & voice-to-text ambient scribe."
        ),
        (
            "Appointments Center",
            "Lightweight booking modal, automated simulated SMS & Email alerts, encounter closure modal with status tracking.",
            "Integration with Twilio Programmable SMS API and SendGrid production Email webhooks."
        ),
        (
            "Executive Dashboard",
            "5 KPI cards (patients impacted, hours saved, cost avoidance), SVG Value Realized and AI Precision charts.",
            "Automated weekly PDF executive report generator & BI export (Tableau/QuickSight)."
        ),
        (
            "Audit & Compliance",
            "RBAC route guards (Admin exclusive), real-time access logs, PHI redaction event stream, security score.",
            "Immutable blockchain audit ledger & AWS CloudTrail SIEM integration."
        ),
        (
            "Backend Infrastructure",
            "Local Vite / Express Node server (`server.ts`) running synthetic FHIR JSON schemas.",
            "AWS Cloud infrastructure (Amplify, Bedrock, HealthLake, ECS Fargate, Aurora Serverless)."
        )
    ]

    for module, supp, fut in matrix:
        row_cells = t_feat.add_row().cells
        row_cells[0].text = module
        row_cells[1].text = f"✅ {supp}"
        row_cells[2].text = f"⏳ {fut}"
        set_cell_bg(row_cells[0], "F8FAFC")

    # Save
    doc_path = "/Users/rk/Antigravity/demo/demo-clinical-ai/Clinical_AI_Portal_Demo_QA_Guide.docx"
    doc.save(doc_path)
    print(f"Successfully created Q&A Word document at: {doc_path}")

if __name__ == "__main__":
    create_demo_qa_docx()
