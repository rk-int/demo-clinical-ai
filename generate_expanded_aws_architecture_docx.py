import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_expanded_aws_architecture_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Color Palette
    PRIMARY_COLOR = RGBColor(15, 23, 42)    # Slate 900 (#0f172a)
    AWS_ORANGE = RGBColor(255, 153, 0)      # AWS Orange (#ff9900)
    SECONDARY_COLOR = RGBColor(2, 132, 199) # Sky 600 (#0284c7)
    EMERALD_COLOR = RGBColor(16, 185, 129)  # Emerald 500 (#10b981)
    PURPLE_COLOR = RGBColor(124, 58, 237)   # Purple 600 (#7c3aed)
    DARK_COLOR = RGBColor(51, 65, 85)       # Slate 700 (#334155)

    def set_cell_bg(cell, hex_color):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("ENTERPRISE HEALTHCARE AI PLATFORM ON AWS")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("HIPAA-Aligned Agentic RAG Reference Architecture — Multi-AZ / Multi-Region Blueprint")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 1: Executive Overview & Comparison Analysis
    h1 = doc.add_heading("1. Architectural Comparison & Gap Analysis", level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    p1 = doc.add_paragraph(
        "Based on a comprehensive review of the Enterprise Healthcare AI Reference Architecture diagram, "
        "we have enriched our AWS production blueprint to cover 10 core architectural domains. "
        "The updated blueprint expands beyond basic LLM and database hosting to include specialized voice AI services, "
        "graph databases, threat discovery, streaming integrations, LLMOps evaluation harnesses, and multi-region disaster recovery."
    )
    p1.runs[0].font.size = Pt(10.5)

    # Section 2: Comprehensive 10-Layer AWS Tech Stack Mapping Table
    h2 = doc.add_heading("2. Comprehensive Feature & AWS Tech Stack Mapping Table", level=1)
    h2.runs[0].font.color.rgb = PRIMARY_COLOR

    t_map = doc.add_table(rows=1, cols=4)
    t_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_map.rows[0].cells
    m_hdr[0].text = "Architecture Domain"
    m_hdr[1].text = "Current Local Prototype"
    m_hdr[2].text = "Target AWS Tech Stack (Full Reference Architecture)"
    m_hdr[3].text = "Architectural Purpose & Technical Role"

    for cell in m_hdr:
        set_cell_bg(cell, "0F172A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    master_mapping = [
        # Domain 1: Edge & Ingress
        (
            "1. User Channels, Secure Edge & Gateway",
            "React 18 SPA, Vite dev server (`npm run dev`), simulated TLS badge",
            "AWS Amplify, Amazon CloudFront, Amazon Route 53, AWS WAF, AWS Shield Advanced, Amazon API Gateway",
            "Global low-latency edge delivery via CloudFront CDN, Route 53 DNS failover, AWS WAF/Shield L7 DDoS protection, API Gateway zero-trust mTLS ingress with rate limiting."
        ),
        # Domain 2: Voice & Multimodal AI
        (
            "2. Clinical Voice & Ambient AI Services",
            "Text-only prompt input in `KnowledgeQAView.tsx`",
            "Amazon Transcribe Medical, Amazon Lex, Amazon Textract, Comprehend Medical, Amazon Polly, Amazon Translate",
            "Speech-to-text for clinical voice dictation (Transcribe Medical), conversational voice bots (Lex), medical OCR chart extraction (Textract), and multi-lingual translation."
        ),
        # Domain 3: Authentication & Identity
        (
            "3. Identity, Access & SSO",
            "Simulated JWT state (`auth.types.ts`) with clinician/admin roles",
            "Amazon Cognito, AWS IAM / IAM Identity Center, AWS Managed SSO",
            "User Pools for clinician login, SAML 2.0 / OIDC federation for hospital EMR SSO (Epic/Cerner IDP), custom JWT role claims (`isAdmin`), and fine-grained IAM policy scoping."
        ),
        # Domain 4: Agentic Orchestration
        (
            "4. Agentic Orchestration & LLM Gateway",
            "State machine state (`VerticalPatientSearchFlowCanvas.tsx`, `agentEngine.ts`)",
            "Amazon Bedrock Agents, AWS Step Functions, AWS Lambda, AWS ECS (Fargate) / EKS",
            "Bedrock Agents + Step Functions orchestrates Orchestrator, Clinical, Workflow, and Compliance agents. Lambda executes serverless tool actions (FHIR queries, orders)."
        ),
        # Domain 5: Context & Session Cache
        (
            "5. Context, Session & Semantic Cache",
            "React in-memory component state",
            "Amazon ElastiCache for Redis, Amazon DynamoDB, Apache Pinot",
            "ElastiCache Redis handles LLM response semantic caching to cut latency & API cost. DynamoDB stores conversation state & active session context. Pinot runs real-time session analytics."
        ),
        # Domain 6: Foundation Models & Guardrails
        (
            "6. Foundation Model Fabric & Guardrails",
            "Simulated LLM response generator in `agentEngine.ts`",
            "Amazon Bedrock (Claude 3.5 Sonnet), SageMaker Endpoints, Amazon Bedrock Guardrails",
            "Bedrock managed access to Anthropic Claude 3.5 Sonnet. SageMaker for fine-tuned clinical models. Bedrock Guardrails for DLP PHI redaction, prompt injection defense, & zero-hallucination checks."
        ),
        # Domain 7: Advanced RAG & Graph Data Stores
        (
            "7. Advanced RAG & Knowledge Graph",
            "Static guideline arrays in `KnowledgeQAView.tsx`",
            "Amazon Bedrock Knowledge Bases, Amazon OpenSearch Service, Amazon Neptune (Health KG)",
            "Bedrock Knowledge Bases parses clinical PDFs. OpenSearch Service handles hybrid vector + BM25 search. Amazon Neptune manages Health Knowledge Graph (Care Pathways & Entity Linking)."
        ),
        # Domain 8: Data Platform & EHR Integration
        (
            "8. Governed Data Platform & Integration",
            "Synthetic FHIR R4 JSON schemas (`syntheticFhirData.ts`)",
            "AWS HealthLake, Amazon S3 Lakehouse, AWS Lake Formation, Amazon Aurora (RDS), Amazon Redshift, AWS AppFlow, Amazon MSK, Amazon MQ",
            "AWS HealthLake manages HIPAA FHIR R4 store. S3 Lakehouse + Glue + Athena for analytics. Lake Formation for FGAC governance. AppFlow/MSK/MQ for HL7v2, DICOM, and streaming EHR events."
        ),
        # Domain 9: Observability, FinOps & LLMOps
        (
            "9. Observability, LLMOps & FinOps",
            "Console logs & step state counters (`currentExecutionStep`)",
            "Amazon CloudWatch, AWS X-Ray, AWS CodePipeline, RAG Evaluation Harness, AWS Cost Explorer",
            "X-Ray traces multi-agent latency across microservices. CloudWatch monitors logs & alerts. CodePipeline automates CI/CD. FinOps tracks token budget & unit costs."
        ),
        # Domain 10: Security, Compliance & Disaster Recovery
        (
            "10. Security Overlay & Multi-Region Resilience",
            "Local environment variables and in-memory execution",
            "AWS KMS, AWS CloudHSM, AWS Secrets Manager, Amazon Macie, Amazon GuardDuty, AWS Security Hub, AWS Config, AWS CloudTrail, AWS Global Accelerator, Aurora Global Database, AWS Backup",
            "KMS envelope encryption, CloudHSM FIPS 140-2 L3 key vault, Macie automated PHI discovery in S3, GuardDuty threat detection, CloudTrail immutable audit logs, Global Accelerator + Aurora Global DB for multi-region failover."
        )
    ]

    for domain, local_impl, aws_svc, role_desc in master_mapping:
        row_cells = t_map.add_row().cells
        row_cells[0].text = domain
        row_cells[1].text = local_impl
        row_cells[2].text = aws_svc
        row_cells[3].text = role_desc
        set_cell_bg(row_cells[0], "F8FAFC")
        set_cell_bg(row_cells[2], "FEF3C7") # Amber highlight for AWS Services

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Summary of Key Added Tech Stack Services
    h3 = doc.add_heading("3. Summary of Key Added AWS Services (Post-Diagram Analysis)", level=1)
    h3.runs[0].font.color.rgb = PRIMARY_COLOR

    added_services = [
        ("Amazon Neptune", "Managed Graph Database powering the Health Knowledge Graph (Care Pathways, Clinical Ontologies, SNOMED/ICD-10 Entity Linking)."),
        ("Amazon Transcribe Medical & Amazon Lex", "HIPAA-eligible speech recognition and voice chat interfaces for hands-free clinician voice assistant dictation."),
        ("Amazon ElastiCache for Redis", "In-memory semantic caching layer to store past LLM prompt embeddings and responses, drastically reducing inference latency and API cost."),
        ("Amazon Macie & Amazon GuardDuty", "Automated PHI/PII sensitive data discovery in S3 data lakes and intelligent threat monitoring across container/agent workloads."),
        ("Amazon Textract & Comprehend Medical", "Medical document OCR extraction and clinical entity detection to ingest unstructured physician PDF notes into FHIR format."),
        ("Amazon MSK & Amazon MQ", "Managed Streaming for Apache Kafka and MQ queues to process real-time HL7v2 and DICOM event streams from hospital PACS/EMR."),
        ("Amazon Aurora Global Database & AWS Global Accelerator", "Multi-region active/warm-standby disaster recovery with automated cross-region database replication (RPO 5 min, RTO 15 min)."),
        ("AWS CodePipeline & RAG Evaluation Harness", "Automated CI/CD deployment pipeline with built-in RAG regression gates (Groundedness, Faithfulness, Recall metrics).")
    ]

    for svc_name, desc in added_services:
        p_svc = doc.add_paragraph()
        p_svc.paragraph_format.left_indent = Inches(0.2)
        p_svc.paragraph_format.space_after = Pt(4)

        r_name = p_svc.add_run(f"🔹 {svc_name}: ")
        r_name.font.bold = True
        r_name.font.color.rgb = PURPLE_COLOR

        r_desc = p_svc.add_run(desc)
        r_desc.font.size = Pt(10)
        r_desc.font.color.rgb = DARK_COLOR

    # Save
    doc_path = "/Users/rk/Antigravity/demo/demo-clinical-ai/Clinical_AI_Portal_AWS_Architecture_Blueprint.docx"
    doc.save(doc_path)
    print(f"Successfully created expanded AWS Architecture Word document at: {doc_path}")

if __name__ == "__main__":
    create_expanded_aws_architecture_docx()
